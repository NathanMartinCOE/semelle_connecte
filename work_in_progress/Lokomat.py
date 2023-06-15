import pandas as pd
import numpy as np
import matplotlib.pyplot as plt

from tkinter import Tk
from tkinter.filedialog import askopenfilename

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from semelle_connecte.Tools.ToolsGetStepEvent import GetStepEvent
from semelle_connecte.Tools.ToolsInterpolationGrf import Interpolation

from pyCGM2.Report import normativeDatasets

### ============================== Arguments ===================================================

root = Tk()
root.withdraw()
DataPath = askopenfilename()

while True:
    key = input("Body mass of the patient integer or float in kg:")
    try:
        mass = int(key)
        break
    except:
        print(f"{key} is not a valid data. Exemple: 80")

# DataPath = "C:\\Users\\Nathan\\Desktop\\Wheelchair tests datas\\Recording_230614_082200.txt"
# mass = 80

### =============================== Normative dataset Schwartz 2008 =============================
NormativeDataset = normativeDatasets.NormativeData("Schwartz2008","Free")
HipAngleMean = pd.DataFrame(NormativeDataset.data["HipMoment"]["mean"], columns=["X","Y","Z"])["X"]
HipAngleSd = pd.DataFrame(NormativeDataset.data["HipMoment"]["sd"], columns=["X","Y","Z"])["X"]
HipAngleUpper = (HipAngleMean + HipAngleSd) / 1000
HipAngleLower = (HipAngleMean - HipAngleSd) / 1000
KneeAngleMean = pd.DataFrame(NormativeDataset.data["KneeMoment"]["mean"], columns=["X","Y","Z"])["X"]
KneeAngleSd = pd.DataFrame(NormativeDataset.data["KneeMoment"]["sd"], columns=["X","Y","Z"])["X"]
KneeAngleUpper = (KneeAngleMean + KneeAngleSd) / 1000
KneeAngleLower = (KneeAngleMean - KneeAngleSd) / 1000

### ================================================ Procedures =======================================================
data = pd.read_csv(DataPath, header=2, sep=";")

for side in ["LEFT", "RIGHT"]:
    data[f"GAIT INDEX {side} (0-1000)"] = round(data[f"GAIT INDEX {side} (0-1000)"])
    data[f"GAIT INDEX {side} (0-1000)"][data[f"GAIT INDEX {side} (0-1000)"]<345] = 0
    data[f"GAIT INDEX {side} (0-1000)"][data[f"GAIT INDEX {side} (0-1000)"]>945] = 0

    for joint in ["HIP", "KNEE"]:
        data[f"{joint} JOINT ANGLE {side} [Degrees]"] = np.rad2deg(data[f"{joint} JOINT ANGLE {side} [Radians]"])
        data[f"UPPER {joint} JOINT ANGLE {side} [Degrees]"] = data[f"{joint} JOINT ANGLE {side} [Degrees]"] + 2
        data[f"LOWER {joint} JOINT ANGLE {side} [Degrees]"] = data[f"{joint} JOINT ANGLE {side} [Degrees]"] - 2
        data[f"{joint} JOINT ANGLE DEV. {side} [Degrees]"] = np.rad2deg(data[f"{joint} JOINT ANGLE DEV. {side} [Radians]"])
        data[f"{joint} PATIENT ANGLE {side} [Degrees]"] = data[f"{joint} JOINT ANGLE {side} [Degrees]"] - data[f"{joint} JOINT ANGLE DEV. {side} [Degrees]"]


HeelStrikeLeft, ToeOffLeft = GetStepEvent(data["GAIT INDEX LEFT (0-1000)"])
HeelStrikeRight, ToeOffRight = GetStepEvent(data["GAIT INDEX RIGHT (0-1000)"])

def NormaliseToeOff(HeelStrike, ToeOff):
    """
    As the curves are normalised over 1000 points, the normalised Toe Off indexes must be calculated.

    Args: 
        HeelStrike (list): List of all index off Heel Strike
        ToeOff (list): List of all index off Toe Off
    Outputs: 
        ToeOff_norm (list): List of all index off Toe Off normalise
    """
    ToeOff_norm = []
    for i in np.arange(len(HeelStrike)-1):
        ToeOff_norm.append((ToeOff[i]-HeelStrike[i])*1000/(HeelStrike[i+1]-HeelStrike[i]))
    return ToeOff_norm

ToeOffLeft_norm = NormaliseToeOff(HeelStrikeLeft, ToeOffLeft)
ToeOffRight_norm = NormaliseToeOff(HeelStrikeRight, ToeOffRight)

def NormaliseStep(data, HeelStrike):
    """ 
    Function for normalising data in % of cycle, by taking a step between two HeelStrikes.

    Args:
        data (pd.DataSeries): data to normalised
        HeelStrike (list): List of all index off Heel Strike
    Outputs:
        mean data in % of cycle
    """
    row = []
    for i in np.arange(len(HeelStrike)-1):
        xnew, ynew = Interpolation(data[HeelStrike[i]:HeelStrike[i+1]],1000)
        row.append(ynew)
    return pd.DataFrame(row).mean(axis=0)

data_normalised = pd.DataFrame()
for joint in ["HIP", "KNEE"]:
    for side, HeelStrike in zip(["LEFT","RIGHT"] , [HeelStrikeLeft, HeelStrikeRight]):
        data_normalised[f"{joint} PATIENT ANGLE {side} [Degrees]"] = NormaliseStep(data[f"{joint} PATIENT ANGLE {side} [Degrees]"], HeelStrike)
        data_normalised[f"{joint} JOINT ANGLE DEV. {side} [Degrees]"] = NormaliseStep(data[f"{joint} JOINT ANGLE DEV. {side} [Degrees]"], HeelStrike)
        data_normalised[f"{joint} JOINT ANGLE {side} [Degrees]"] = NormaliseStep(data[f"{joint} JOINT ANGLE {side} [Degrees]"], HeelStrike)
        data_normalised[f"UPPER {joint} JOINT ANGLE {side} [Degrees]"] = NormaliseStep(data[f"UPPER {joint} JOINT ANGLE {side} [Degrees]"], HeelStrike)
        data_normalised[f"LOWER {joint} JOINT ANGLE {side} [Degrees]"] = NormaliseStep(data[f"LOWER {joint} JOINT ANGLE {side} [Degrees]"], HeelStrike)
        data_normalised[f"{joint} JOINT TORQUE {side} [NM]"] = NormaliseStep(data[f"{joint} JOINT TORQUE {side} [NM]"], HeelStrike) / mass

x, data_normalised["UPPER HIP JOINT ANGLE [NM]"] = Interpolation(HipAngleUpper, 1000)
x, data_normalised["LOWER HIP JOINT ANGLE [NM]"] = Interpolation(HipAngleLower, 1000)
x, data_normalised["UPPER KNEE JOINT ANGLE [NM]"] = Interpolation(KneeAngleUpper, 1000)
x, data_normalised["LOWER KNEE JOINT ANGLE [NM]"] = Interpolation(KneeAngleLower, 1000)

### =============================== Report (based on data_normalised DataFrame) ==================================

document = Document()
document.add_heading('Lokomat repport', 0)

### ================= BWS =======================================
document.add_heading('Body weight support during session', 1)
plt.figure()
plt.plot(data["TARGET BWS [KG]"], label="Target BWS (kg)", c="red")
plt.plot(data["ACTUAL BWS [KG]"], label="Actual BWS (kg)", c="blue")
plt.hlines(y=mass, xmin=0, xmax=len(data["ACTUAL BWS [KG]"]), label="Patient mass (kg)", colors="black")
plt.legend()
plt.savefig('LokomatRepport.png')
plt.clf()
document.add_picture('LokomatRepport.png', height = Inches(3.5))

### ================= Kinematics =======================================
document.add_heading('Kinematics during session', 1)

plt.figure()
plt.title("Hip")
plt.plot(data_normalised["HIP PATIENT ANGLE LEFT [Degrees]"], label="Left", c="r")
plt.plot(data_normalised["HIP PATIENT ANGLE RIGHT [Degrees]"], label="Right", c="b")
plt.vlines(x= np.mean(ToeOffLeft_norm), ymin=-20, ymax=70, colors="r")
plt.vlines(x= np.mean(ToeOffRight_norm), ymin=-20, ymax=70, colors="b")
plt.fill_between(x=range(0,1000), y1=data_normalised["LOWER HIP JOINT ANGLE LEFT [Degrees]"], y2=data_normalised["UPPER HIP JOINT ANGLE LEFT [Degrees]"], color='lightpink')
plt.fill_between(x=range(0,1000), y1=data_normalised["LOWER HIP JOINT ANGLE RIGHT [Degrees]"], y2=data_normalised["UPPER HIP JOINT ANGLE RIGHT [Degrees]"], color='lightblue')
plt.ylim((-20,70))
plt.savefig('LokomatRepport.png')
plt.clf()
document.add_picture('LokomatRepport.png', height = Inches(3.5))

plt.figure()
plt.title("Knee")
plt.plot(data_normalised["KNEE PATIENT ANGLE LEFT [Degrees]"], label="Left", c="r")
plt.plot(data_normalised["KNEE PATIENT ANGLE RIGHT [Degrees]"], label="Right", c="b")
plt.vlines(x= np.mean(ToeOffLeft_norm), ymin=-15, ymax=75, colors="r")
plt.vlines(x= np.mean(ToeOffRight_norm), ymin=-15, ymax=75, colors="b")
plt.fill_between(x=range(0,1000), y1=data_normalised["LOWER KNEE JOINT ANGLE LEFT [Degrees]"], y2=data_normalised["UPPER KNEE JOINT ANGLE LEFT [Degrees]"], color='lightpink')
plt.fill_between(x=range(0,1000), y1=data_normalised["LOWER KNEE JOINT ANGLE RIGHT [Degrees]"], y2=data_normalised["UPPER KNEE JOINT ANGLE RIGHT [Degrees]"], color='lightblue')
plt.ylim((-15,75))
plt.savefig('LokomatRepport.png')
plt.clf()
document.add_picture('LokomatRepport.png', height = Inches(3.5))

# ========================= Kinetics ================================================
document.add_heading('Kinetics during session', 1)

plt.figure()
plt.title("Hip")
plt.plot(data_normalised["HIP JOINT TORQUE LEFT [NM]"], label="Left", c="r")
plt.plot(data_normalised["HIP JOINT TORQUE RIGHT [NM]"], label="Right", c="b")
plt.vlines(x= np.mean(ToeOffLeft_norm), ymin=-20, ymax=70, colors="r")
plt.vlines(x= np.mean(ToeOffRight_norm), ymin=-20, ymax=70, colors="b")
plt.fill_between(x=range(0,1000),y1=data_normalised["LOWER HIP JOINT ANGLE [NM]"], y2=data_normalised["UPPER HIP JOINT ANGLE [NM]"], color="lightgrey")
plt.ylim((-2,3))
plt.savefig('LokomatRepport.png')
plt.clf()
document.add_picture('LokomatRepport.png', height = Inches(3.5))

plt.figure()
plt.title("Knee")
plt.plot(data_normalised["KNEE JOINT TORQUE LEFT [NM]"], label="Left", c="r")
plt.plot(data_normalised["KNEE JOINT TORQUE RIGHT [NM]"], label="Right", c="b")
plt.vlines(x= np.mean(ToeOffLeft_norm), ymin=-15, ymax=75, colors="r")
plt.vlines(x= np.mean(ToeOffRight_norm), ymin=-15, ymax=75, colors="b")
plt.fill_between(x=range(0,1000),y1=data_normalised["LOWER KNEE JOINT ANGLE [NM]"], y2=data_normalised["UPPER KNEE JOINT ANGLE [NM]"], color="lightgrey")
plt.ylim((-1,1))
plt.savefig('LokomatRepport.png')
plt.clf()
document.add_picture('LokomatRepport.png', height = Inches(3.5))

document.save("C:\\Users\\Nathan\\Desktop\\Wheelchair tests datas\\Lokomat_Report.docx")