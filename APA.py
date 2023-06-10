# test APA
import os
import pandas as pd
import matplotlib.pyplot as plt


from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ALIGN_VERTICAL

from Walking.Walking import Walking
from SOLE.FeetMe import readFeetMeMultipleCsv
from SOLE.FeetMe import ReadSpatioTemporalCsv
from Reader.Reader import Reader
from Tools.ToolsSymetryIndex import SymetryIndex

ID = 161095
mass = 60
tests = ["Parcours_1", "Parcours_2", "Haie", "Salom", "doubleTache", "bande", "Back"]

document = Document()

for test in tests:

    document.add_heading(f'Condition : {test}', 0)
    
    DataPath = f"C:\\Users\\Nathan\\Desktop\\Wheelchair tests datas\\FeetMe\\APA\\{ID}\\{test}\\"

    files = os.listdir(DataPath)
    fullfilename_walking = []
    try :
        walking_name = [file for file in files if file.endswith(".hdf5")][0]
        fullfilename_walking = os.path.join(DataPath, walking_name)
    except:
        pass
    metrics_file = [file for file in files if file.endswith("metrics.csv")][0]
    fullfilename_metrics = os.path.join(DataPath, metrics_file)
    pressures_files = [file for file in files if file.endswith("pressure.csv")]
    fullfilenames_pressures = []
    for pressures_file in pressures_files:
        fullfilenames_pressures.append(os.path.join(DataPath, pressures_file))


    try:
        walking = Reader(fullfilename_walking).readh5()
        
    except:
        SoleInstanceRight, SoleInstanceLeft = readFeetMeMultipleCsv(fullfilenames = fullfilenames_pressures, freq = 110)
        DataFrameSpatioTemporal_Left, DataFrameSpatioTemporal_Right = ReadSpatioTemporalCsv(fullfilename_metrics)
        
        walking = Walking(mass)
        walking.setLeftLegSole(SoleInstanceLeft)
        walking.setRightLegSole(SoleInstanceRight)
        walking.setDataFrameSpatioTemporal_Left(DataFrameSpatioTemporal_Left)
        walking.setDataFrameSpatioTemporal_Right(DataFrameSpatioTemporal_Right)

        from Walking.WalkingFilters import WalkingKinematicsFilter
        from Walking.WalkingKinematicsProcedure import GroundReactionForceKinematicsProcedure
        procedure = GroundReactionForceKinematicsProcedure()
        WalkingKinematicsFilter(walking, procedure).run()

        from Walking.WalkingFilters import WalkingDataProcessingFilter
        from Walking.WalkingDataProcessingProcedure import DeleteStepProcedure
        procedure = DeleteStepProcedure()
        WalkingDataProcessingFilter(walking, procedure).run()

        from Walking.WalkingFilters import WalkingDataProcessingFilter
        from Walking.WalkingDataProcessingProcedure import NormalisationProcedure
        procedure = NormalisationProcedure()
        WalkingDataProcessingFilter(walking, procedure).run()

        from Writer.Writer import Writer
        StoragePathHDF5 = DataPath
        Writer(walking = walking, path = StoragePathHDF5, file_name = f"walking_{ID}_{test}.hdf5").writeh5()

    document.add_heading('Kinetic gait parameters', level=1)

    ### Assymetry of the Vertical Ground Reaction Force
    from Walking.WalkingFilters import WalkingGraphicsFilter
    from Walking.WalkingGraphicsProcedure import PlotDynamicSymetryFunctionNormalisedProcedure
    procedure = PlotDynamicSymetryFunctionNormalisedProcedure(show_graph = False, save_graph = True, StoragePath = DataPath)
    WalkingGraphicsFilter(walking, procedure).run()

    files = os.listdir(DataPath)
    figs = [file for file in files if file.endswith("VerticalGrf.png")]
    for fig in figs:
        FigPath = os.path.join(DataPath, fig)
        document.add_paragraph('Assymetry of the Vertical Ground Reaction Force').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_picture(FigPath, height = Inches(3.5))


    ### Evolution of Vertical Ground Reaction Force during the test'
    from Walking.WalkingFilters import WalkingDataProcessingFilter
    from Walking.WalkingDataProcessingProcedure import CutDataProcessingProcedure
    procedure = CutDataProcessingProcedure()
    procedure.setCutNumber(n_cut=3)
    WalkingDataProcessingFilter(walking, procedure).run()

    from Walking.WalkingFilters import WalkingGraphicsFilter
    from Walking.WalkingGraphicsProcedure import PlotCutGroundReactionForceProcedure
    procedure = PlotCutGroundReactionForceProcedure(show_graph = False, save_graph = True, StoragePath = DataPath)
    WalkingGraphicsFilter(walking, procedure).run()
    
    files = os.listdir(DataPath)
    figs = [file for file in files if file.endswith("CutDataGrf.png")]
    for fig in figs:
        FigPath = os.path.join(DataPath, fig)
        document.add_paragraph('Evolution of Vertical Ground Reaction Force during the test').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_picture(FigPath, width = Inches(7.5))


    ### ================================================ Spatio-Temporal ================================================
    document.add_heading('Spatio-temporal gait parameters', level=1)

    p = document.add_paragraph("This test contains:")
    if walking.m_DataFrameSpatioTemporal_Left.shape[0] == 0:
        p.add_run(" No data for Left Leg")
    elif walking.m_DataFrameSpatioTemporal_Left.shape[0] > 0:
        p.add_run(f" {walking.m_DataFrameSpatioTemporal_Left.shape[0]} steps for Left Leg")
    p.add_run(" and")
    if walking.m_DataFrameSpatioTemporal_Right.shape[0] == 0:
        p.add_run(" No data for Right Leg")
    elif walking.m_DataFrameSpatioTemporal_Right.shape[0] > 0:
        p.add_run(f" {walking.m_DataFrameSpatioTemporal_Right.shape[0]} steps for Left Leg")
    
    ### Evolution of Spatio-temporal parameters
    from Walking.WalkingFilters import WalkingGraphicsFilter
    from Walking.WalkingGraphicsProcedure import PlotSpatioTemporalParametersEvolutionProcedure
    procedure = PlotSpatioTemporalParametersEvolutionProcedure(show_graph = False, save_graph = True, StoragePath = DataPath)
    WalkingGraphicsFilter(walking, procedure).run()

    files = os.listdir(DataPath)
    figs = [file for file in files if file.endswith("SpatioTemporalParametersEvolution.png")]
    for fig in figs:
        FigPath = os.path.join(DataPath, fig)
        document.add_paragraph('Evolution of Spatio-temporal parameters during the test').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_picture(FigPath, width = Inches(5.9))

    ### Comparaison of Spatio-temporal parameters Asymetry
    from Walking.WalkingFilters import WalkingGraphicsFilter
    from Walking.WalkingGraphicsProcedure import PlotSpatioTemporalParametersBoxplotProcedure
    procedure = PlotSpatioTemporalParametersBoxplotProcedure(show_graph = False, save_graph = True, StoragePath = DataPath)
    WalkingGraphicsFilter(walking, procedure).run()

    files = os.listdir(DataPath)
    figs = [file for file in files if file.endswith("SpatioTemporalParametersBoxplot.png")]
    for fig in figs:
        FigPath = os.path.join(DataPath, fig)
        document.add_paragraph('Boxplot of Spatio-temporal parameters by legs').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_picture(FigPath, width = Inches(5.9))

        
    ### ================================================ Table ================================================
    document.add_heading('Summary table for kinetic an spatio-temporal parameters', level=1)

    n_round = 2

    SummaryTable = pd.DataFrame()
    SummaryTable["Parameters"] = ["Vertical Ground Reaction Force",
                                  "Single Support Duration (ms)", 
                                  "Double Support Duration (ms)", 
                                  "Stance Duration (ms)", 
                                  "Swing Duration (ms)"]
    SummaryTable["Left_mean"] = [round(walking.m_sole["LeftLeg"].data["VerticalGrf"].mean(), n_round),
                                round(walking.m_DataFrameSpatioTemporal_Left["singleSupportDuration (ms)"].mean(), n_round),
                                round(walking.m_DataFrameSpatioTemporal_Left["doubleSupportDuration (ms)"].mean(), n_round),
                                round(walking.m_DataFrameSpatioTemporal_Left["stanceDuration (ms)"].mean(), n_round),
                                round(walking.m_DataFrameSpatioTemporal_Left["swingDuration (ms)"].mean(), n_round)]
    SummaryTable["Left"] = [f'{round(walking.m_sole["LeftLeg"].data["VerticalGrf"].mean(), n_round)} ± {round(walking.m_sole["LeftLeg"].data["VerticalGrf"].std(), n_round)}',
    f'{round(walking.m_DataFrameSpatioTemporal_Left["singleSupportDuration (ms)"].mean(), n_round)} ± {round(walking.m_DataFrameSpatioTemporal_Left["singleSupportDuration (ms)"].std(), n_round)}',
    f'{round(walking.m_DataFrameSpatioTemporal_Left["doubleSupportDuration (ms)"].mean(), n_round)} ± {round(walking.m_DataFrameSpatioTemporal_Left["doubleSupportDuration (ms)"].std(), n_round)}',
    f'{round(walking.m_DataFrameSpatioTemporal_Left["stanceDuration (ms)"].mean(), n_round)} ± {round(walking.m_DataFrameSpatioTemporal_Left["stanceDuration (ms)"].std(), n_round)}',
    f'{round(walking.m_DataFrameSpatioTemporal_Left["swingDuration (ms)"].mean(), n_round)} ± {round(walking.m_DataFrameSpatioTemporal_Left["swingDuration (ms)"].std(), n_round)}']
    SummaryTable["Right_mean"] = [round(walking.m_sole["RightLeg"].data["VerticalGrf"].mean(), n_round),
                                round(walking.m_DataFrameSpatioTemporal_Right["singleSupportDuration (ms)"].mean(), n_round),
                                round(walking.m_DataFrameSpatioTemporal_Right["doubleSupportDuration (ms)"].mean(), n_round),
                                round(walking.m_DataFrameSpatioTemporal_Right["stanceDuration (ms)"].mean(), n_round),
                                round(walking.m_DataFrameSpatioTemporal_Right["swingDuration (ms)"].mean(), n_round)]
    SummaryTable["Right"] = [f'{round(walking.m_sole["RightLeg"].data["VerticalGrf"].mean(), n_round)} ± {round(walking.m_sole["RightLeg"].data["VerticalGrf"].std(), n_round)}',
    f'{round(walking.m_DataFrameSpatioTemporal_Right["singleSupportDuration (ms)"].mean(), n_round)} ± {round(walking.m_DataFrameSpatioTemporal_Right["singleSupportDuration (ms)"].std(), n_round)}',
    f'{round(walking.m_DataFrameSpatioTemporal_Right["doubleSupportDuration (ms)"].mean(), n_round)} ± {round(walking.m_DataFrameSpatioTemporal_Right["doubleSupportDuration (ms)"].std(), n_round)}',
    f'{round(walking.m_DataFrameSpatioTemporal_Right["stanceDuration (ms)"].mean(), n_round)} ± {round(walking.m_DataFrameSpatioTemporal_Right["stanceDuration (ms)"].std(), n_round)}',
    f'{round(walking.m_DataFrameSpatioTemporal_Right["swingDuration (ms)"].mean(), n_round)} ± {round(walking.m_DataFrameSpatioTemporal_Right["swingDuration (ms)"].std(), n_round)}'] 
    SummaryTable["Symmetry Index"] = SymetryIndex(SummaryTable["Left_mean"], SummaryTable["Right_mean"])
    SummaryTable = SummaryTable.drop(columns=["Left_mean", "Right_mean"])

    table = document.add_table(rows=SummaryTable.shape[0] + 1, cols=SummaryTable.shape[1])
    table.style = 'Table Grid'
    for i, column_name in enumerate(SummaryTable.columns):
        table.cell(0, i).text = column_name
    for i, row in SummaryTable.iterrows():
        for j, value in enumerate(row):
            table.cell(i + 1, j).text = str(value)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_VERTICAL.CENTER

    document.add_paragraph("Symetry Index is represented as a percentage where perfect symmetry between legs is 0%.").italic=True
    try:
        FigPath = "C:\\Users\\Nathan\\Desktop\\Recherche\\Github\\semelle_connecte\\ImageForReport\\SI_Formule.jpg"
        document.add_picture(FigPath, width = Inches(1.57))
    except:
        print("Don't find ImageForReport ----> SI_Formule.jpg")
    
document.save(f"C:\\Users\\Nathan\\Desktop\\Wheelchair tests datas\\FeetMe\\APA\\{ID}\\Report.docx")
import ipdb; ipdb.set_trace()