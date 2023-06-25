"""
Coding : utf8
Author : Nathan Martin
Create : 30/05/2023

=========================================================  Documentation =========================================================

This file is used to process all walking_{name}_test{X}.hdf5 files for compute the standard error of measurement.

Input :
    PathMSE: the path to the folder containing all the hdf5 files (walking).

hdf5 file (walking):
    The files are created by the python script (Feetme_SEM_Create_Walking_hdf5.py)
    File contents :
        walking having been instantiated with the reaction forces given by the Feetme soles (readFeetmeCsv or readFeetmeMultipleCsv)
        3 procedures: GroundReactionForceKinematicsProcedure   (for more informations please read procedure documentation)
                       StandardisationProcedure                (for more informations please read procedure documentation)
                       DynamicSymetryFunctionComputeProcedure  (for more informations please read procedure documentation)
        All stored in the .hdf5 file
    File storage:
        walking_{name}_test{X}.hdf5 // in PathMSE

Output:
------> A pdf file with :
--> Plots for each condition with 2 sublopts for each leg : 
    Left plot for left leg // Right plot for right leg
    For each subplot the SEM for the leg is write in the title.
    Legend :
        Red   : mean ground reaction force of Left leg 
        Blue  : mean ground reaction force of Right leg
        Black : standard error of measurement at each time
        Dashed gray : all ground reaction force for the corresponding leg
--> Table of SEM for each condition for each metrics 
------> A csv file with all data of metrics for stat in R
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import os
import re

from matplotlib.backends.backend_pdf import PdfPages
from semelle_connecte.Reader.Reader import Reader
from semelle_connecte.Tools.RichardBacker_SEM import RichardBacker_SEM


def main():
        
    pd.options.mode.chained_assignment = None # Caution this disable the warning of indexing pandas


    conditions = ["normal_ground", "mottek_vL12_vR12", "mottek_vL12_vR14", "mottek_vL12_vR16", "mottek_vL12_vR18", "mottek_vL14_vR12", "mottek_vL16_vR12", "mottek_vL18_vR12"]

    pdf = PdfPages("C:\\Users\\Nathan\\Desktop\\Wheelchair tests datas\\FeetMe\\MSE\\Graphiques_SEM_GRF.pdf")
    row_data_result_SEM = []
    row_data_result_value = []
    all_data_metric = []

    for condition in conditions:
        PathMSE = f"C:\\Users\\Nathan\\Desktop\\Wheelchair tests datas\\FeetMe\\MSE\\{condition}\\"
        
        files = os.listdir(PathMSE)
        csv_files = [file for file in files if file.endswith(".csv")]
        hdf5_files = [file for file in files if file.endswith(".hdf5")]

        ### =============== Pressure ==================
        frames = ["Frame_" + str(i) for i in range(1000)]
        columns_names_pressure = ["id", "test", "leg"] + frames
        data_pressure = []

        for file in hdf5_files:
            PathHDF5 = PathMSE
            NameFileHDF5 = str(file)
            DataPathHDF5 = os.path.join(PathHDF5, NameFileHDF5)
            walking = Reader(DataPathHDF5).readh5()

            ### ======== look for the name of participant ===========
            find_name = re.search(r"walking_(.*?)_test", NameFileHDF5)
            if find_name:
                name = find_name.group(1)
            else:
                print("No name find. ----------- Please be sure to name the hdf5 as walking_name_testX.hdf5 ")
            ### ======== look for the number of the test session =====
            find_N_test = re.search(r"test(.*?).hdf5", NameFileHDF5)
            if find_N_test:
                N_test = find_N_test.group(1)
            else:
                print("Number of test not find. ----------- Please be sure to name the hdf5 as walking_name_testX.hdf5 ")

            ### == Instanciated one row_data witch contain "name", "test number", "leg side", GRF value for 0-1000 =====
            DataMeanGrf = pd.DataFrame()
            for leg in ["LeftLeg", "RightLeg"]:
                DataStepGrfValue = pd.DataFrame()
                for step in np.arange(0, len(walking.m_StepGrfValue[leg]["VerticalGrf"])):
                    DataStepGrfValue[f"step{step}"] = walking.m_StepGrfValue[leg]["VerticalGrf"][step]
                DataMeanGrf[leg] = DataStepGrfValue.mean(axis=1, skipna=True)
                row_data = [name, N_test, leg]
                row_data = row_data + DataMeanGrf[leg].values.tolist()
                data_pressure.append(row_data)

        Score_SEM_Left, Score_SEM_Right = RichardBacker_SEM(data = data_pressure,
                                                            columns_names = columns_names_pressure, 
                                                            condition = condition, 
                                                            pdf = pdf, 
                                                            plot_graph = True)
        
        row_data_result_SEM.append([condition, "Left", "Vertical Ground Reaction", round(Score_SEM_Left,4)])
        row_data_result_SEM.append([condition, "Right", "Vertical Ground Reaction", round(Score_SEM_Right,4)])

        ### =============== Metric ==================
        metrics = ["stanceDuration (ms)", "singleSupportDuration (ms)", "doubleSupportDuration (ms)", "swingDuration (ms)"]
        for metric in metrics:

            ### =============== Metric ==================
            columns_names_metric = ["id", "test", "leg", metric]
            data_metric = []
                
            for file in csv_files:
                Path_csv = PathMSE
                NameFile_csv = str(file)
                DataPath_csv = os.path.join(Path_csv, NameFile_csv)

                ### ======== look for the number of the test session =====
                find_N_test = re.search(r"test(.*?).csv", NameFile_csv)
                if find_N_test:
                    N_test = find_N_test.group(1)
                else:
                    print("Number of test not find. ----------- Please be sure to name the csv as testX.csv ")

                Feetme_metrics = pd.read_csv(DataPath_csv, header=1)
                Feetme_metric = Feetme_metrics.loc[:, ["side", metric]]
                Feetme_metric["side"][Feetme_metric["side"] == "left"] = "LeftLeg"
                Feetme_metric["side"][Feetme_metric["side"] == "right"] = "RightLeg"

                ### == Instanciated one row_data witch contain "name", "test number", "leg side", metric =====
                for side in ["LeftLeg", "RightLeg"]:
                    row_data = [name, N_test, side]
                    row_data = row_data + Feetme_metric[Feetme_metric["side"] == side].mean(axis=0, skipna=True).values.tolist()
                    ### ============ Add row_data to data for the DataFrame DataRaw ===========================
                    data_metric.append(row_data)
                    row_data_result_value.append([condition, side, metric, Feetme_metric[Feetme_metric["side"] == side].mean(axis=0, skipna=True).values[0], Feetme_metric[Feetme_metric["side"] == side].std(axis=0, skipna=True).values[0]])
                    ### ============ Add value of each step for all metric ====================================
                    row_data_all_metric = [name, N_test, condition, side, metric]
                    row_data_all_metric = row_data_all_metric + Feetme_metric[metric][Feetme_metric["side"] == side].values.tolist()
                    all_data_metric.append(row_data_all_metric)
                    
            Score_SEM_Left, Score_SEM_Right = RichardBacker_SEM(data = data_metric, 
                                                                columns_names = columns_names_metric, 
                                                                condition = condition, 
                                                                pdf = pdf, 
                                                                plot_graph = False)
            
            row_data_result_SEM.append([condition, "Left", metric, round(Score_SEM_Left,4)])
            row_data_result_SEM.append([condition, "Right", metric, round(Score_SEM_Right,4)])
        
            
    ### ================================ Raw Result SEM ==========================================

    Result_SEM = pd.DataFrame(row_data_result_SEM, columns=["Condition", "Leg", "Parametre", "SEM"])

    ### ================================ Prety table SEM ==========================================
    Prety_Result_SEM = pd.DataFrame()

    Prety_Result_SEM["Parametre"] = Result_SEM[Result_SEM["Condition"]== "normal_ground"]["Parametre"].values
    Prety_Result_SEM["Side"] = ["Left", "Right"] * 5
    Prety_Result_SEM["normal_ground"]    = Result_SEM[Result_SEM["Condition"]== "normal_ground"]["SEM"].values
    Prety_Result_SEM["mottek_vL12_vR12"] = Result_SEM[Result_SEM["Condition"]== "mottek_vL12_vR12"]["SEM"].values
    Prety_Result_SEM["mottek_vL12_vR14"] = Result_SEM[Result_SEM["Condition"]== "mottek_vL12_vR14"]["SEM"].values
    Prety_Result_SEM["mottek_vL12_vR16"] = Result_SEM[Result_SEM["Condition"]== "mottek_vL12_vR16"]["SEM"].values
    Prety_Result_SEM["mottek_vL12_vR18"] = Result_SEM[Result_SEM["Condition"]== "mottek_vL12_vR18"]["SEM"].values
    Prety_Result_SEM["mottek_vL14_vR12"] = Result_SEM[Result_SEM["Condition"]== "mottek_vL14_vR12"]["SEM"].values
    Prety_Result_SEM["mottek_vL16_vR12"] = Result_SEM[Result_SEM["Condition"]== "mottek_vL16_vR12"]["SEM"].values
    Prety_Result_SEM["mottek_vL18_vR12"] = Result_SEM[Result_SEM["Condition"]== "mottek_vL18_vR12"]["SEM"].values

    fig_Prety_Result_SEM, ax = plt.subplots(figsize=(12,4))
    ax.axis('tight')
    ax.axis('off')
    the_table = ax.table(cellText=Prety_Result_SEM.values,colLabels=Prety_Result_SEM.columns,loc='center')
    pdf.savefig(fig_Prety_Result_SEM, bbox_inches='tight')

    pdf.close() ### Close pdf file

    ### =============================== Save DatFrame with all metrics in a csv ======================================

    max_length = max(len(values) for values in all_data_metric)                                     # Find max step  
    Steps = ["Step_" + str(i) for i in range(max_length-5)]                                         # Create col names  
    columns_names_all_metric = ["Name", "N_test", "Condition", "Leg", "Metric"] + Steps             # Create col names
    DatFrame_all_metric = pd.DataFrame(data = all_data_metric, columns = columns_names_all_metric)  # Create DataFrame

    PathSaveData = "C:\\Users\\Nathan\\Desktop\\Wheelchair tests datas\\FeetMe\\MSE\\"
    DatFrame_all_metric.to_csv(os.path.join(PathSaveData, "DatFrame_all_metric.csv"))



from docx import Document
from docx.shared import Inches

def Graph3D():
    conditions = ["normal_ground", "mottek_vL12_vR12", "mottek_vL12_vR14", "mottek_vL12_vR16", "mottek_vL12_vR18", "mottek_vL14_vR12", "mottek_vL16_vR12", "mottek_vL18_vR12"]

    document = Document()
    StoragePath = "C:\\Users\\Nathan\\Desktop\\Wheelchair tests datas\\FeetMe\\MSE\\img\\"

    for condition in conditions:
        document.add_heading(f'Condition : {condition}', 0)
        PathMSE = f"C:\\Users\\Nathan\\Desktop\\Wheelchair tests datas\\FeetMe\\MSE\\{condition}\\"
        
        files = os.listdir(PathMSE)
        hdf5_files = [file for file in files if file.endswith(".hdf5")]

        for file in hdf5_files:
            PathHDF5 = PathMSE
            NameFileHDF5 = str(file)
            DataPathHDF5 = os.path.join(PathHDF5, NameFileHDF5)
            walking = Reader(DataPathHDF5).readh5()

            from semelle_connecte.Walking.WalkingFilters import WalkingKinematicsFilter
            from semelle_connecte.Walking.WalkingKinematicsProcedure import GroundReactionForceKinematicsProcedure
            procedure = GroundReactionForceKinematicsProcedure()
            WalkingKinematicsFilter(walking, procedure).run()

            from semelle_connecte.Walking.WalkingFilters import WalkingDataProcessingFilter
            from semelle_connecte.Walking.WalkingDataProcessingProcedure import NormalisationProcedure
            procedure = NormalisationProcedure()
            WalkingDataProcessingFilter(walking, procedure).run()

            from semelle_connecte.Walking.WalkingFilters import WalkingGraphicsFilter
            from semelle_connecte.Walking.WalkingGraphicsProcedure import PlotVerticalGroundReaction3DProcedure

            procedure = PlotVerticalGroundReaction3DProcedure(show_graph = False,  save_graph = True, StoragePath = StoragePath)
            WalkingGraphicsFilter(walking, procedure).run()

            document.add_picture(os.path.join(StoragePath,'VerticalGroundReaction3D.png'), height = Inches(3.5))
    
    document.save("C:\\Users\\Nathan\\Desktop\\Wheelchair tests datas\\FeetMe\\MSE\\Graphiques3D_GRF.docx")




if __name__ == '__main__':
    # main()
    Graph3D()